# Importar bibliotecas necesarias
from moviepy.editor import VideoFileClip, AudioFileClip
from google.cloud import speech_v1p1beta1 as speech, translate_v2 as translate, vision
from google.oauth2 import service_account
from docx import Document
import io
import os
import math
import csv

# Configuración de credenciales de Google Cloud
credentials = service_account.Credentials.from_service_account_file('./keys/summarize-415522-ce93d8d9fb9a.json')
client_speech = speech.SpeechClient(credentials=credentials)
client_translate = translate.Client(credentials=credentials)
client_vision = vision.ImageAnnotatorClient(credentials=credentials)

# Función para extraer audio del video
def extraer_audio(video_path):
    audio_path = 'temp_audio.wav'
    if not os.path.exists(audio_path):
        print("Extrayendo audio...")
        video = VideoFileClip(video_path)
        video.audio.write_audiofile(audio_path)
    else:
        print("El archivo de audio ya existe, omitiendo extracción.")
    return audio_path

# Modificación aquí: Duración de segmento ajustada a 45 segundos
def transcribir_audio_por_partes(audio_path, duracion_segmento=45, csv_path='transcripciones.csv'):
    clip_audio = AudioFileClip(audio_path)
    duracion_total = clip_audio.duration
    transcript_total = ""
    
    # Intenta abrir el archivo CSV para leer los segmentos ya transcritos (si existe)
    segmentos_transcritos = {}
    if os.path.exists(csv_path):
        with open(csv_path, mode='r', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                segmento_index, transcripcion = row
                segmentos_transcritos[segmento_index] = transcripcion
    
    # Asegurarse de que cada segmento de audio tiene un nombre único
    segmento_index = 0
    with open(csv_path, mode='a', encoding='utf-8', newline='') as csvfile:
        writer = csv.writer(csvfile)
        
        for inicio in range(0, math.ceil(duracion_total), duracion_segmento):
            fin = min(inicio + duracion_segmento, duracion_total)
            segmento_audio = clip_audio.subclip(inicio, fin)
            segmento_path = f"temp_segment_{segmento_index}.wav"
            segmento_index_str = str(segmento_index)
            
            # Si el segmento ya ha sido transcrito, usar esa transcripción
            if segmento_index_str in segmentos_transcritos:
                transcript_segmento = segmentos_transcritos[segmento_index_str]
            else:
                segmento_audio.write_audiofile(segmento_path, codec='pcm_s16le', ffmpeg_params=["-ar", "16000", "-ac", "1"])
                transcript_segmento = transcribir_segmento(segmento_path)
                writer.writerow([segmento_index_str, transcript_segmento])  # Guardar transcripción en el CSV
                csvfile.flush()  # Forzar la escritura en el archivo inmediatamente
                os.remove(segmento_path)  # Limpiar el archivo de audio temporal
            
            transcript_total += " " + transcript_segmento
            segmento_index += 1
    
    return transcript_total.strip()


# Ajuste en la función de transcripción para usar la configuración modificada
def transcribir_segmento(segmento_path):
    with open(segmento_path, 'rb') as audio_file:
        content = audio_file.read()
    audio = speech.RecognitionAudio(content=content)
    config = speech.RecognitionConfig(
        encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
        sample_rate_hertz=16000,  # Ajuste a la tasa de muestreo correcta
        language_code='en-US',  # Ajusta según necesidad
        enable_automatic_punctuation=True
    )

    try:
        response = client_speech.recognize(config=config, audio=audio)
        return ' '.join([result.alternatives[0].transcript for result in response.results])
    except Exception as e:
        print(f"Error al transcribir el segmento: {segmento_path}, error: {str(e)}")
        return ""

# Función modificada para traducir texto en fragmentos
def traducir_texto(texto, target_language='es'):
    # Definir el tamaño máximo del fragmento en caracteres
    # Ajusta este valor para estar seguro de no exceder el límite de la API
    max_fragmento_size = 10000  # Aproximadamente 10 KB por fragmento
    fragmentos = [texto[i:i+max_fragmento_size] for i in range(0, len(texto), max_fragmento_size)]
    
    texto_traducido = ""
    for fragmento in fragmentos:
        result = client_translate.translate(fragmento, target_language=target_language)
        texto_traducido += result['translatedText'] + " "
    
    return texto_traducido.strip()


# Función para analizar imágenes y obtener descripciones
def analizar_imagenes(imagen_path):
    with io.open(imagen_path, 'rb') as image_file:
        content = image_file.read()
    image = vision.Image(content=content)
    response = client_vision.label_detection(image=image)
    descripciones = [label.description for label in response.label_annotations]
    return ', '.join(descripciones)

# Función para extraer y analizar fotogramas clave del video
def extraer_y_analizar_fotogramas(video_path, intervalo=60):
    clip = VideoFileClip(video_path)
    duracion = int(clip.duration)
    descripciones = []
    for i in range(0, duracion, intervalo):
        frame_path = f"temp_frame_{i}.jpg"
        clip.save_frame(frame_path, t=i)
        descripcion = analizar_imagenes(frame_path)
        descripciones.append((i, descripcion))
        os.remove(frame_path)
    return descripciones

# Función para crear el documento DOCX
def crear_documento(texto, descripciones, file_name='Resumen_Video.docx'):
    document = Document()
    document.add_heading('Resumen del Video', level=0)
    document.add_paragraph(texto)
    document.add_heading('Descripción de Imágenes', level=1)
    for tiempo, descripcion in descripciones:
        document.add_paragraph(f"Tiempo: {tiempo}s - Descripción: {descripcion}")
    document.save(file_name)

# Integración del proceso completo
def procesar_video(video_path):
    print("Extrayendo audio...")
    audio_path = extraer_audio(video_path)
    
    print("Transcribiendo audio a texto...")
    texto = transcribir_audio_por_partes(audio_path)
    
    print("Traduciendo texto a español...")
    texto_es = traducir_texto(texto)
    
    print("Extrayendo y analizando fotogramas clave...")
    descripciones = extraer_y_analizar_fotogramas(video_path)
    
    print("Creando documento DOCX con el resumen y descripciones de imágenes...")
    crear_documento(texto_es, descripciones)

    print("Proceso completado. El documento DOCX ha sido creado.")

    # Limpieza del archivo de audio temporal
    os.remove(audio_path)

# Ejecutar el proceso para un video específico
video_path = './videos/course.mp4'
procesar_video(video_path)
